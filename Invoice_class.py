from PyQt5.QtWidgets import *
from workenv import create_workdir
from docx import Document
import os


class Invoice(QInputDialog):
    """This class represents an invoice, containing its specific attributes
     and the methods that apply to it"""
    def __init__(self, parent=None):
        super(Invoice, self).__init__(parent)
        self.doc = Document()

    def make_invoice(self):
        """Prompt the user to specify a name for the invoice, create it and save
        it in the working directory """
        work_dir = create_workdir()
        inv_input = QInputDialog()
        text, ok = inv_input.getText(self, "Invoice name",
                                     "Please enter a name for the invoice:               ")
        if ok and text != '':
            self.name = text
            paragraph = self.doc.add_paragraph('Invoice Report', 'Title')
            paragraph.add_run("\nLorem ipsum")
            os.chdir(work_dir)
            self.doc.save(f"{self.name}" + ".docx")
        else:
            pass

    def set_status(self):
        pass

    def get_status(self):
        pass

