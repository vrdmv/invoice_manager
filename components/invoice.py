from PyQt5.QtWidgets import *
from workenv import create_workdir
from docx import Document
import os
from database import *


class Invoice(QInputDialog):
    """This class represents an invoice, containing its specific attributes
     and the methods that apply to it"""
    def __init__(self, parent=None):
        super().__init__(parent)

    def make_invoice(self):
        """Prompt the user to specify a name for the invoice, create it and save
        it in the working directory """
        self.doc = Document()
        work_dir = create_workdir()
        inv_input = QInputDialog()
        active = True
        while active:
            text, ok = inv_input.getText(self, "Invoice name",
                                     "Please enter a name for the invoice:               ")
            if ok and text != '':
                self.name = text
                paragraph = self.doc.add_paragraph('Invoice Report', 'Title')
                paragraph.add_run("\nLorem ipsum")
                path = f"{work_dir}" + "\\" + f"{self.name}" + ".docx"
                if not os.path.exists(path):
                    os.chdir(work_dir)
                    self.doc.save(f"{self.name}" + ".docx")
                    set_initstatus(self.name)
                    active = False
                else:
                    self.show_popup()
            else:
                active = False

    def show_popup(self):
        """Show a popup message if invoice name already exists"""
        msg = QMessageBox()
        msg.setWindowTitle("File/folder already exists.")
        msg.setText("The invoice or project you are trying to create "
                    "already exists."
                    " Please specify another name.")
        msg.setIcon(QMessageBox.Warning)
        x = msg.exec_()
