from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
import sys
import shutil
import send2trash
import os
import docx

class Ui_MainWindow(QWidget):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(1120, 800)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.frame = QtWidgets.QFrame(self.centralwidget)
        self.frame.setGeometry(QtCore.QRect(10, 10, 1091, 711))
        self.frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame.setObjectName("frame")

        # New Project button characteristics & function call
        self.project_button = QtWidgets.QPushButton(self.frame)
        self.project_button.setGeometry(QtCore.QRect(10, 80, 121, 41))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(9)
        self.project_button.setFont(font)
        self.project_button.setObjectName("project_button")
        self.project_button.clicked.connect(self.make_project)

        # New Invoice button characteristics & function call
        self.invoice_button = QtWidgets.QPushButton(self.frame)
        self.invoice_button.setGeometry(QtCore.QRect(10, 20, 121, 41))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(9)
        self.invoice_button.setFont(font)
        self.invoice_button.setObjectName("invoice_button")
        self.invoice_button.clicked.connect(self.make_invoice)

        # ListView model characteristics & functionality
        self.listView = QtWidgets.QListView(self.frame)
        self.listView.setGeometry(QtCore.QRect(10, 190, 221, 471))
        self.listView.setAlternatingRowColors(True)
        self.listView.setAcceptDrops(True)
        self.listView.setDropIndicatorShown(True)
        self.listView.setDragEnabled(True)
        self.listView.setSelectionMode(QAbstractItemView.SingleSelection)
        self.listView.setMovement(QListView.Snap)
        self.listView.setDefaultDropAction(Qt.MoveAction)
        self.listView.setObjectName("listView")
        self.listView.setEditTriggers(QAbstractItemView.NoEditTriggers)

        path = r"C:\Invoice Manager"

        self.dirModel = QtWidgets.QFileSystemModel()
        self.dirModel.setRootPath(QtCore.QDir.rootPath())
        # self.dirModel.setFilter(QtCore.QDir.NoDotAndDotDot | QtCore.QDir.AllDirs)
        self.dirModel.setReadOnly(False)

        self.listView.setModel(self.dirModel)
        self.listView.setRootIndex(self.dirModel.index(path))

        # TreeView model characteristics & functionality
        self.fileModel = QtWidgets.QFileSystemModel()
        # self.fileModel.setFilter(QtCore.QDir.NoDotAndDotDot | QtCore.QDir.Files)
        self.fileModel.setReadOnly(False)
        
        self.treeView = QtWidgets.QTreeView(self.frame)
        self.treeView.setModel(self.fileModel)
        self.treeView.setRootIndex(self.fileModel.index(path))
        self.treeView.setDragEnabled(True)
        self.treeView.setAcceptDrops(True)
        self.treeView.setDropIndicatorShown(True)
        self.treeView.setSortingEnabled(True)
        self.treeView.setColumnWidth(0, 400)
        self.treeView.setColumnHidden(1, True)
        self.treeView.setGeometry(QtCore.QRect(270, 90, 801, 571))
        self.treeView.setObjectName("treeView")
        self.treeView.setSelectionMode(QAbstractItemView.ContiguousSelection)
        self.treeView.setDefaultDropAction(Qt.MoveAction)
        self.treeView.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.treeView.setContextMenuPolicy(Qt.CustomContextMenu)
        
        # Radio button characterisics
        self.radioButton = QtWidgets.QRadioButton(self.frame)
        self.radioButton.setGeometry(QtCore.QRect(460, 60, 95, 20))
        self.radioButton.setObjectName("Draft")
        self.radioButton_2 = QtWidgets.QRadioButton(self.frame)
        self.radioButton_2.setGeometry(QtCore.QRect(570, 60, 95, 20))
        self.radioButton_2.setObjectName("Approved")
        self.radioButton_3 = QtWidgets.QRadioButton(self.frame)
        self.radioButton_3.setGeometry(QtCore.QRect(690, 60, 95, 20))
        self.radioButton_3.setObjectName("Dispatched")
        self.radioButton_4 = QtWidgets.QRadioButton(self.frame)
        self.radioButton_4.setGeometry(QtCore.QRect(810, 60, 95, 20))
        self.radioButton_4.setObjectName("Paid")
        self.radioButton_5 = QtWidgets.QRadioButton(self.frame)
        self.radioButton_5.setGeometry(QtCore.QRect(930, 60, 95, 20))
        self.radioButton_5.setObjectName("Overdue")

        # Status bar characteristis
        self.progressBar = QtWidgets.QProgressBar(self.frame)
        self.progressBar.setGeometry(QtCore.QRect(270, 670, 801, 23))
        self.progressBar.setToolTip("Invoice status")
        self.progressBar.setAutoFillBackground(False)
        self.progressBar.setStyleSheet("")
        self.progressBar.setMaximum(100)
        self.progressBar.setTextVisible(False)
        self.progressBar.setObjectName("progressBar")

        # When folder in listView is clicked, the items stored in it are
        # displayed in the treeview.
        self.listView.clicked.connect(self.on_click)

        # Open item in treeview when double-clicked
        self.treeView.doubleClicked.connect(self.on_dblclick)

        # Open context menu when item clicked in treeView.
        self.treeView.customContextMenuRequested.connect(self.openMenu)
        
        # Update progress bar based on radio button toggled
        self.radioButton.toggled['bool'].connect(lambda: self.progressBar.setValue(20))
        self.radioButton_2.toggled['bool'].connect(lambda: self.progressBar.setValue(40))
        self.radioButton_3.toggled['bool'].connect(lambda: self.progressBar.setValue(60))
        self.radioButton_4.toggled['bool'].connect(lambda: self.progressBar.setValue(80))
        self.radioButton_5.toggled['bool'].connect(lambda: self.progressBar.setValue(100))

        # Labels
        self.project_label = QtWidgets.QLabel(self.frame)
        self.project_label.setGeometry(QtCore.QRect(90, 160, 61, 21))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(10)
        self.project_label.setFont(font)
        self.project_label.setObjectName("project_label")
        self.invoice_label = QtWidgets.QLabel(self.frame)
        self.invoice_label.setGeometry(QtCore.QRect(270, 60, 81, 21))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(10)
        self.invoice_label.setFont(font)
        self.invoice_label.setObjectName("invoice_label")

        # MainWindow characteristics
        MainWindow.setCentralWidget(self.centralwidget)
        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)
        

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.project_button.setText(_translate("MainWindow", "+ New Project"))
        self.project_button.setShortcut(_translate("MainWindow", "Ctrl+P"))
        self.project_button.setToolTip("Shortcut: Ctrl+P")
        self.invoice_button.setText(_translate("MainWindow", "+ New Invoice"))
        self.invoice_button.setShortcut(_translate("MainWindow", "Ctrl+N"))
        self.invoice_button.setToolTip("Shortcut: Ctrl+N")
        self.project_label.setText(_translate("MainWindow", "Projects"))
        self.invoice_label.setText(_translate("MainWindow", "Invoices"))
        self.radioButton.setText(_translate("MainWindow", "Draft"))
        self.radioButton_2.setText(_translate("MainWindow", "Approved"))
        self.radioButton_3.setText(_translate("MainWindow", "Dispatched"))
        self.radioButton_4.setText(_translate("MainWindow", "Paid"))
        self.radioButton_5.setText(_translate("MainWindow", "Overdue"))


    # Program Functions

    def on_click(self, index):
        """When a directory is selected in the list view, show all of its files
        in the tree view."""
        path = self.dirModel.fileInfo(index).absoluteFilePath()
        self.treeView.setRootIndex(self.fileModel.setRootPath(path))


    def on_dblclick(self, index):
        """ Open file upon double-click."""
        path = self.fileModel.fileInfo(index).absoluteFilePath()
        path_str = str(path)
        os.startfile(path_str)


    def dlt(self, index):
        """ Send the selected file to the recycle bin"""
        path = self.fileModel.fileInfo(index).absoluteFilePath()
        # file_name = self.fileModel.fileName(index)
        # file_name_str = str(file_name)
        send2trash.send2trash(os.path.abspath(path))


    def move_to_archive(self, index):
        """Move selected file to archive"""
        source = self.fileModel.fileInfo(index).absoluteFilePath()
        destination = create_archive()
        #file_name = self.fileModel.fileName(index)
        #file_info = self.fileModel.fileInfo(index)
        shutil.move(os.path.abspath(source), os.path.abspath(destination))


    def copy(self, index):
        """Copy the selected file within the same directory."""
        source = self.fileModel.fileInfo(index).absoluteFilePath()
        copy_str = "(copy)"
        destination = self.dirModel.fileInfo(index).absoluteFilePath()
        destination_str = str(destination)
        destination_firstfew = destination_str[:-5]
        destintion_final4 = destination_str[-5:]
        final_destination = destination_firstfew + f"{copy_str}" + destintion_final4
        shutil.copy(os.path.abspath(source), os.path.abspath(final_destination))


    def openMenu(self, position):
        """Setup a context menu, containing various options, to open upon right
        click."""
        index = self.treeView.indexAt(position)
        if not index.isValid():
            return
        menu = QMenu(self)
        rename_action = menu.addAction("Rename")
        copy_action = menu.addAction("Copy")
        delete_action = menu.addAction("Send to trashbin")
        archive_action = menu.addAction("Move to archive")
        action = menu.exec_(self.treeView.viewport().mapToGlobal(position))

        try:
            if action == delete_action:
                self.dlt(index)
            if action == rename_action:
                self.treeView.edit(index)
            if action == copy_action:
                self.copy(index)
            if action == archive_action:
                self.move_to_archive(index)

        except PermissionError:
            print("Nope, can't do that!")


    def make_invoice(self):
        """Prompt the user to specify a name for the invoice,  create the invoice
        and save in the working directory."""
        working_directory = create_workdir()
        inv_inputbox = QInputDialog
        text, ok = inv_inputbox.getText(self, "Invoice name",
                                        "Please enter invoice name: ",
                                        QLineEdit.Normal)
        if ok and text != '':
            invoice_name = text
            invoice = docx.Document()
            invoice.add_paragraph('Invoice Report', 'Title')
            paraobj1 = invoice.add_paragraph('This is the second paragraph.')
            paraobj2 = invoice.add_paragraph('This is yet another paragraph.')
            paraobj1.add_run('This text is being added to '
                             'the second paragraph.')
            os.chdir(working_directory)
            invoice.save(f"{invoice_name}" + ".docx")
        else:
            pass


    def make_project(self):
        """Create a new project folder."""
        active = True
        while active:
            text, ok = QInputDialog.getText(self, "QInputDialog.getText",
                                            "Please enter project name: ",
                                            QLineEdit.Normal)
            if ok and text != '':
                project_name = text
                project_dir = os.path.join("C:\\", "Invoice Manager",
                                           f"{project_name}")

                if not os.path.exists(project_dir):
                    os.makedirs(project_dir)
                    active = False
            else:
                continue



def create_workdir():
    """Create a current working directory.
    (works across all operating systems)"""

    work_dir = os.path.join("C:\\", "Invoice Manager",
                            "Work environment")
    if not os.path.exists(work_dir):
        os.makedirs(work_dir)
    work_dir_str = str(work_dir)
    return work_dir_str


create_workdir()


def create_archive():
    """Create a current working directory.
    (works across all operating systems)"""

    archive_dir = os.path.join("C:\\", "Invoice Manager",
                               "Archive")
    if not os.path.exists(archive_dir):
        os.makedirs(archive_dir)
    archive_dir_str = str(archive_dir)
    return archive_dir_str


create_archive()


if __name__ == "__main__":
    app = QApplication(sys.argv)

    app.setStyle('Fusion')
    dark_palette = QPalette()
    dark_palette.setColor(QPalette.Window, QColor(53, 53, 53))
    dark_palette.setColor(QPalette.WindowText, Qt.white)
    dark_palette.setColor(QPalette.Base, QColor(25, 25, 25))
    dark_palette.setColor(QPalette.AlternateBase, QColor(53, 53, 53))
    dark_palette.setColor(QPalette.ToolTipBase, Qt.white)
    dark_palette.setColor(QPalette.ToolTipText, Qt.white)
    dark_palette.setColor(QPalette.Text, Qt.white)
    dark_palette.setColor(QPalette.Button, QColor(53, 53, 53))
    dark_palette.setColor(QPalette.ButtonText, Qt.white)
    dark_palette.setColor(QPalette.BrightText, Qt.red)
    dark_palette.setColor(QPalette.Link, QColor(42, 130, 218))
    dark_palette.setColor(QPalette.Highlight, QColor(42, 130, 218))
    dark_palette.setColor(QPalette.HighlightedText, Qt.black)
    qApp.setPalette(dark_palette)
    qApp.setStyleSheet(
        "QToolTip { color: #ffffff; background-color: #2a82da;"
        " border: 1px solid white; }")

    MainWindow = QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())
