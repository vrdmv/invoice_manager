from PyQt5.QtWidgets import *
from invoiceman_gui import Ui_MainWindow
from Palette_class import dark_palette
from create_workenv import *
import sys
import shutil
import send2trash
import os
import docx

class Logic(QMainWindow, Ui_MainWindow):
    def __init__(self):
        super(Logic, self).__init__()

        # Set up the user interface
        self.setupUi(self)

        # "New Project" button function call
        self.project_button.clicked.connect(self.make_project)

        # "New Invoice" button function call
        self.invoice_button.clicked.connect(self.make_invoice)

        # When folder in listView is clicked, the items stored in it are
        # displayed in the treeview.
        self.listView.clicked.connect(self.on_click)

        # Open item in treeview when double-clicked
        self.treeView.doubleClicked.connect(self.on_dblclick)

        # Open context menu when item clicked in treeView.
        self.treeView.customContextMenuRequested.connect(self.openMenu)

        # Update progress bar based on radio button toggled
        self.radioButton.toggled['bool'].connect(
            lambda: self.progressBar.setValue(20))
        self.radioButton_2.toggled['bool'].connect(
            lambda: self.progressBar.setValue(40))
        self.radioButton_3.toggled['bool'].connect(
            lambda: self.progressBar.setValue(60))
        self.radioButton_4.toggled['bool'].connect(
            lambda: self.progressBar.setValue(80))
        self.radioButton_5.toggled['bool'].connect(
            lambda: self.progressBar.setValue(100))

    # Program Functions
    def on_click(self, index):
        """When a directory is selected in the list view, show all of its files
        in the tree view."""
        path = self.dirModel.fileInfo(index).absoluteFilePath()
        self.treeView.setRootIndex(self.fileModel.setRootPath(path))


    def on_dblclick(self, index):
        """ Open file upon double-click."""
        path = self.fileModel.fileInfo(index).absoluteFilePath()
        path_str = str(path)
        os.startfile(path_str)


    def dlt(self, index):
        """ Send the selected file to the recycle bin"""
        path = self.fileModel.fileInfo(index).absoluteFilePath()
        # file_name = self.fileModel.fileName(index)
        # file_name_str = str(file_name)
        send2trash.send2trash(os.path.abspath(path))


    def move_to_archive(self, index):
        """Move selected file to archive"""
        source = self.fileModel.fileInfo(index).absoluteFilePath()
        destination = create_archive()
        # file_name = self.fileModel.fileName(index)
        # file_info = self.fileModel.fileInfo(index)
        shutil.move(os.path.abspath(source), os.path.abspath(destination))


    def copy(self, index):
        """Copy the selected file within the same directory."""
        source = self.fileModel.fileInfo(index).absoluteFilePath()
        copy_str = "(copy)"
        destination = self.dirModel.fileInfo(index).absoluteFilePath()
        destination_str = str(destination)
        destination_firstfew = destination_str[:-5]
        destintion_final4 = destination_str[-5:]
        final_destination = destination_firstfew + f"{copy_str}" + destintion_final4
        shutil.copy(os.path.abspath(source),
                    os.path.abspath(final_destination))


    def openMenu(self, position):
        """Setup a context menu, containing various options, to open upon right
        click."""
        index = self.treeView.indexAt(position)
        if not index.isValid():
            return
        menu = QMenu(self)
        rename_action = menu.addAction("Rename")
        copy_action = menu.addAction("Copy")
        delete_action = menu.addAction("Send to trashbin")
        archive_action = menu.addAction("Move to archive")
        action = menu.exec_(self.treeView.viewport().mapToGlobal(position))
        try:
            if action == delete_action:
                self.dlt(index)
            if action == rename_action:
                self.treeView.edit(index)
            if action == copy_action:
                self.copy(index)
            if action == archive_action:
                self.move_to_archive(index)
        except PermissionError:
            print("Nope, can't do that!")


    def make_invoice(self):
        """Prompt the user to specify a name for the invoice, create the invoice
        and save in the working directory."""
        working_directory = create_workdir()
        inv_inputbox = QInputDialog()
        text, ok = inv_inputbox.getText(self, "Invoice name",
                                        "Please enter invoice name:                       ",
                                        QLineEdit.Normal)
        if ok and text != '':
            invoice_name = text
            invoice = docx.Document()
            invoice.add_paragraph('Invoice Report', 'Title')
            paraobj1 = invoice.add_paragraph(
                'This is the second paragraph.')
            paraobj2 = invoice.add_paragraph(
                'This is yet another paragraph.')
            paraobj1.add_run('This text is being added to '
                             'the second paragraph.')
            os.chdir(working_directory)
            invoice.save(f"{invoice_name}" + ".docx")
        else:
            pass


    def make_project(self):
        """Create a new project folder."""
        active = True
        while active:
            text, ok = QInputDialog.getText(self, "Project name",
                                            "Please enter project name:                    ",
                                            QLineEdit.Normal)
            if ok and text != '':
                project_name = text
                project_dir = os.path.join("C:\\", "Invoice Manager",
                                           f"{project_name}")
                if not os.path.exists(project_dir):
                    os.makedirs(project_dir)
                    active = False
            else:
                active = False

create_workdir()
create_archive()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    qApp.setPalette(dark_palette)
    qApp.setStyleSheet("QToolTip { color: #ffffff; background-color: #2a82da;"
        " border: 1px solid white; }")
    logic = Logic()
    logic.show()
    sys.exit(app.exec_())
